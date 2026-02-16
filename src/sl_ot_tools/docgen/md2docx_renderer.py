"""
md2docx_renderer.py — Converts markdown-it-py tokens into a python-docx Document.

Supports: headings, bold/italic/strikethrough, tables, bullet/numbered lists,
code blocks, inline code, blockquotes, horizontal rules, links, footnotes,
and LaTeX math (rendered as styled text).
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt

try:
    from mdit_py_plugins.dollarmath import dollarmath_plugin
    from mdit_py_plugins.footnote import footnote_plugin
    HAS_PLUGINS = True
except ImportError:
    HAS_PLUGINS = False


def create_parser() -> MarkdownIt:
    """Create a markdown-it parser with extensions enabled."""
    md = MarkdownIt("commonmark")
    md.enable("table")
    md.enable("strikethrough")
    if HAS_PLUGINS:
        dollarmath_plugin(md)
        footnote_plugin(md)
    return md


class DocxRenderer:
    """Walks markdown-it tokens and builds a python-docx Document."""

    def __init__(self, author: str = "", title: str = ""):
        self.doc = Document()
        self.author = author
        self.title = title
        self._extracted_title = ""
        self._setup_styles()

        # State
        self._current_paragraph = None
        self._inline_state = []         # stack: 'strong', 'em', 'strikethrough'
        self._list_stack = []           # stack of ('bullet'|'ordered', depth)
        self._pending_list_style = None
        self._in_blockquote = False

        # Table buffering
        self._table_state = None        # None, 'head', 'body'
        self._table_buffer = None
        self._current_row_cells = None

        # Footnotes
        self._footnotes = {}            # id -> list of inline tokens
        self._footnote_counter = 0
        self._footnote_map = {}         # footnote label -> number

    def _setup_styles(self):
        """Configure base document styles."""
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

    # ===================================================================
    # Public API
    # ===================================================================

    def render(self, tokens: list) -> Document:
        """Process all tokens and return the Document."""
        # First pass: collect footnote definitions
        self._collect_footnotes(tokens)

        # Second pass: render
        i = 0
        while i < len(tokens):
            token = tokens[i]
            handler = getattr(self, f"_handle_{token.type}", None)
            if handler:
                handler(token)
            i += 1

        # Append footnotes section
        self._render_footnotes()

        # Set metadata
        self._set_metadata()

        return self.doc

    # ===================================================================
    # Footnote collection (first pass)
    # ===================================================================

    def _collect_footnotes(self, tokens):
        """Scan tokens for footnote definitions and store their content."""
        i = 0
        while i < len(tokens):
            tok = tokens[i]
            if tok.type == "footnote_open":
                fn_id = tok.meta.get("id", 0)
                fn_label = tok.meta.get("label", str(fn_id))
                self._footnote_counter += 1
                self._footnote_map[fn_label] = self._footnote_counter
                # Collect tokens until footnote_close
                content_tokens = []
                i += 1
                depth = 1
                while i < len(tokens) and depth > 0:
                    if tokens[i].type == "footnote_open":
                        depth += 1
                    elif tokens[i].type == "footnote_close":
                        depth -= 1
                        if depth == 0:
                            break
                    else:
                        content_tokens.append(tokens[i])
                    i += 1
                self._footnotes[self._footnote_counter] = content_tokens
            i += 1
        # Reset counter for rendering pass
        self._footnote_counter = 0

    # ===================================================================
    # Block-level handlers
    # ===================================================================

    def _handle_heading_open(self, token):
        level = int(token.tag[1])  # h1 -> 1, h2 -> 2, etc.
        self._current_paragraph = self.doc.add_heading("", level=level)
        # Capture first heading as document title
        if not self._extracted_title and level == 1:
            self._extracted_title = "__pending__"

    def _handle_heading_close(self, token):
        self._current_paragraph = None

    def _handle_paragraph_open(self, token):
        if self._table_state:
            return  # paragraphs inside table cells handled differently
        if self._pending_list_style:
            self._current_paragraph = self.doc.add_paragraph(
                style=self._pending_list_style
            )
            self._pending_list_style = None
        elif self._in_blockquote:
            para = self.doc.add_paragraph()
            pf = para.paragraph_format
            pf.left_indent = Cm(1.5)
            # Add left border via XML
            pPr = para.paragraph_format.element.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:space"), "4")
            left.set(qn("w:color"), "CCCCCC")
            pBdr.append(left)
            pPr.append(pBdr)
            self._current_paragraph = para
        else:
            self._current_paragraph = self.doc.add_paragraph()

    def _handle_paragraph_close(self, token):
        if self._table_state:
            return
        self._current_paragraph = None

    def _handle_bullet_list_open(self, token):
        depth = len(self._list_stack)
        self._list_stack.append(("bullet", depth))

    def _handle_bullet_list_close(self, token):
        if self._list_stack:
            self._list_stack.pop()

    def _handle_ordered_list_open(self, token):
        depth = len(self._list_stack)
        self._list_stack.append(("ordered", depth))

    def _handle_ordered_list_close(self, token):
        if self._list_stack:
            self._list_stack.pop()

    def _handle_list_item_open(self, token):
        if not self._list_stack:
            return
        list_type, depth = self._list_stack[-1]
        if list_type == "bullet":
            style_map = {0: "List Bullet", 1: "List Bullet 2", 2: "List Bullet 3"}
        else:
            style_map = {0: "List Number", 1: "List Number 2", 2: "List Number 3"}
        self._pending_list_style = style_map.get(depth, style_map[2])

    def _handle_list_item_close(self, token):
        pass

    def _handle_fence(self, token):
        """Fenced code block."""
        para = self.doc.add_paragraph()
        run = para.add_run(token.content.rstrip("\n"))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
        # Gray background shading
        pPr = para.paragraph_format.element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F6F8FA")
        shd.set(qn("w:val"), "clear")
        pPr.append(shd)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(4)

    def _handle_code_block(self, token):
        self._handle_fence(token)

    def _handle_hr(self, token):
        """Horizontal rule."""
        para = self.doc.add_paragraph()
        pPr = para.paragraph_format.element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "AAAAAA")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _handle_blockquote_open(self, token):
        self._in_blockquote = True

    def _handle_blockquote_close(self, token):
        self._in_blockquote = False

    def _handle_math_block(self, token):
        """Block-level LaTeX equation — rendered as styled text."""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(token.content.strip())
        run.font.name = "Cambria Math"
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Footnote block handlers (skip during render, already collected) ---
    def _handle_footnote_open(self, token):
        pass

    def _handle_footnote_close(self, token):
        pass

    def _handle_footnote_block_open(self, token):
        pass

    def _handle_footnote_block_close(self, token):
        pass

    def _handle_footnote_anchor(self, token):
        pass

    # ===================================================================
    # Table handlers (buffered approach)
    # ===================================================================

    def _handle_table_open(self, token):
        self._table_buffer = {"headers": [], "rows": [], "col_count": 0}
        self._table_state = "init"

    def _handle_thead_open(self, token):
        self._table_state = "head"

    def _handle_thead_close(self, token):
        self._table_state = "body"

    def _handle_tbody_open(self, token):
        self._table_state = "body"

    def _handle_tbody_close(self, token):
        pass

    def _handle_tr_open(self, token):
        self._current_row_cells = []

    def _handle_tr_close(self, token):
        if self._table_state == "head":
            self._table_buffer["headers"] = self._current_row_cells
            self._table_buffer["col_count"] = len(self._current_row_cells)
        else:
            self._table_buffer["rows"].append(self._current_row_cells)
        self._current_row_cells = None

    def _handle_th_open(self, token):
        self._current_row_cells.append([])

    def _handle_th_close(self, token):
        pass

    def _handle_td_open(self, token):
        self._current_row_cells.append([])

    def _handle_td_close(self, token):
        pass

    def _handle_table_close(self, token):
        """Flush the buffered table into a python-docx table."""
        buf = self._table_buffer
        cols = buf["col_count"]
        if cols == 0:
            self._table_state = None
            self._table_buffer = None
            return

        total_rows = 1 + len(buf["rows"])
        table = self.doc.add_table(rows=total_rows, cols=cols)
        table.style = "Light Grid Accent 1"

        # Header row
        for col_idx, cell_tokens in enumerate(buf["headers"]):
            if col_idx >= cols:
                break
            cell = table.rows[0].cells[col_idx]
            para = cell.paragraphs[0]
            self._current_paragraph = para
            for child in cell_tokens:
                self._handle_inline_child(child)
            for run in para.runs:
                run.bold = True

        # Data rows
        for row_idx, row_cells in enumerate(buf["rows"]):
            for col_idx, cell_tokens in enumerate(row_cells):
                if col_idx >= cols:
                    break
                cell = table.rows[row_idx + 1].cells[col_idx]
                para = cell.paragraphs[0]
                self._current_paragraph = para
                for child in cell_tokens:
                    self._handle_inline_child(child)

        self._table_state = None
        self._table_buffer = None
        self._current_paragraph = None
        # Add spacing after table
        self.doc.add_paragraph("")

    # ===================================================================
    # Inline handler
    # ===================================================================

    def _handle_inline(self, token):
        """Process inline tokens — either into a paragraph or a table cell buffer."""
        if not token.children:
            return

        # If we're inside a table cell, buffer the children
        if self._table_state and self._current_row_cells is not None:
            current_cell = self._current_row_cells[-1] if self._current_row_cells else None
            if current_cell is not None:
                current_cell.extend(token.children)
                return

        # Otherwise render directly into the current paragraph
        if not self._current_paragraph:
            return

        for child in token.children:
            self._handle_inline_child(child)

        # Capture title from first h1
        if self._extracted_title == "__pending__":
            text = self._current_paragraph.text
            if text:
                self._extracted_title = text

    def _handle_inline_child(self, token):
        """Process a single inline child token into the current paragraph."""
        if not self._current_paragraph:
            return

        t = token.type

        if t == "text":
            run = self._current_paragraph.add_run(token.content)
            self._apply_inline_formatting(run)

        elif t == "strong_open":
            self._inline_state.append("strong")
        elif t == "strong_close":
            if "strong" in self._inline_state:
                self._inline_state.remove("strong")

        elif t == "em_open":
            self._inline_state.append("em")
        elif t == "em_close":
            if "em" in self._inline_state:
                self._inline_state.remove("em")

        elif t == "s_open":
            self._inline_state.append("strikethrough")
        elif t == "s_close":
            if "strikethrough" in self._inline_state:
                self._inline_state.remove("strikethrough")

        elif t == "code_inline":
            run = self._current_paragraph.add_run(token.content)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            self._apply_inline_formatting(run)

        elif t == "softbreak":
            self._current_paragraph.add_run("\n")

        elif t == "hardbreak":
            run = self._current_paragraph.add_run()
            run.add_break()

        elif t == "math_inline":
            run = self._current_paragraph.add_run(token.content)
            run.font.name = "Cambria Math"
            run.font.italic = True
            self._apply_inline_formatting(run)

        elif t == "link_open":
            self._inline_state.append("link")

        elif t == "link_close":
            if "link" in self._inline_state:
                self._inline_state.remove("link")

        elif t == "image":
            alt = token.content or token.children[0].content if token.children else "image"
            run = self._current_paragraph.add_run(f"[Image: {alt}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        elif t == "footnote_ref":
            fn_id = token.meta.get("id", 0)
            fn_label = token.meta.get("label", str(fn_id))
            fn_num = self._footnote_map.get(fn_label, "?")
            run = self._current_paragraph.add_run(str(fn_num))
            run.font.superscript = True
            run.font.size = Pt(9)

        elif t == "html_inline":
            # Skip raw HTML
            pass

    def _apply_inline_formatting(self, run):
        """Apply accumulated inline formatting state to a run."""
        if "strong" in self._inline_state:
            run.bold = True
        if "em" in self._inline_state:
            run.italic = True
        if "strikethrough" in self._inline_state:
            run.font.strike = True
        if "link" in self._inline_state:
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True

    # ===================================================================
    # Footnotes rendering
    # ===================================================================

    def _render_footnotes(self):
        """Append footnotes as a 'Notes' section at the end."""
        if not self._footnotes:
            return

        self.doc.add_paragraph()  # spacer
        para = self.doc.add_paragraph()
        run = para.add_run("Notes")
        run.bold = True
        run.font.size = Pt(10)

        for num in sorted(self._footnotes.keys()):
            tokens = self._footnotes[num]
            para = self.doc.add_paragraph()
            run = para.add_run(f"[{num}] ")
            run.font.superscript = True
            run.font.size = Pt(9)

            self._current_paragraph = para
            # Render footnote content tokens
            for tok in tokens:
                if tok.type == "inline" and tok.children:
                    for child in tok.children:
                        self._handle_inline_child(child)
                elif tok.type == "paragraph_open" or tok.type == "paragraph_close":
                    pass  # skip paragraph wrappers inside footnotes

            para.paragraph_format.space_after = Pt(2)
            for r in para.runs:
                if not r.font.size:
                    r.font.size = Pt(9)

        self._current_paragraph = None

    # ===================================================================
    # Metadata
    # ===================================================================

    def _set_metadata(self):
        """Set document core properties."""
        props = self.doc.core_properties
        props.author = self.author
        title = self.title or self._extracted_title
        if title and title != "__pending__":
            props.title = title
        props.created = datetime.now()
        props.modified = datetime.now()


# ===================================================================
# Public API
# ===================================================================

def render_markdown_to_docx(
    md_text: str,
    output_path: str,
    author: str = "",
    title: str = "",
) -> None:
    """Parse markdown text and create a DOCX file.

    Args:
        md_text: Markdown source text.
        output_path: Path to write the .docx file.
        author: Document author name.
        title: Document title (auto-extracted from first H1 if not given).
    """
    parser = create_parser()
    tokens = parser.parse(md_text)
    renderer = DocxRenderer(author=author, title=title)
    doc = renderer.render(tokens)
    doc.save(output_path)
